#!/usr/bin/env python3
"""Generate a DOCX presentation for RecruitAI ATS project."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
import os

SCREENSHOTS_DIR = os.path.join(os.path.dirname(__file__), 'screenshots')
OUTPUT_PATH = os.path.join(os.path.dirname(__file__), 'RecruitAI_Presentation.docx')

doc = Document()

# Page setup - A4 landscape for better screenshot display
for section in doc.sections:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ─── Helper functions ───
def add_title_page(title, subtitle):
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)  # Indigo

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(subtitle)
    run2.font.size = Pt(16)
    run2.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

def add_section_page(title, description):
    doc.add_page_break()
    doc.add_paragraph()
    h = doc.add_heading(title, level=1)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)
    p = doc.add_paragraph(description)
    p.style.font.size = Pt(12)

def add_screenshot(filename, caption):
    img_path = os.path.join(SCREENSHOTS_DIR, filename)
    if os.path.exists(img_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(img_path, width=Inches(9.5))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(caption)
        r.italic = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

def add_bullet_list(items):
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

# ─── Title Page ───
add_title_page('RecruitAI', 'Systeme de Suivi des Candidatures (ATS)\nPropulse par l\'Intelligence Artificielle')

# ─── Project Overview ───
add_section_page(
    'Presentation du Projet',
    'RecruitAI est une plateforme moderne de gestion du recrutement (ATS) '
    'concue pour simplifier et optimiser le processus d\'embauche grace a '
    'l\'intelligence artificielle.'
)
add_bullet_list([
    'Stack technique : Laravel 11 + Inertia.js + Vue 3 + Tailwind CSS',
    'Base de donnees : MySQL avec soft deletes et multi-tenancy',
    'IA : Integration OpenAI pour l\'analyse de CV et le scoring des candidats',
    'Facturation : Stripe via Laravel Cashier (plans Free, Starter, Professional)',
    'Authentification : Laravel Breeze avec RBAC (5 roles)',
    'Interface : Entierement en francais, design responsive',
])

# ─── Architecture ───
add_section_page(
    'Architecture Technique',
    'L\'application suit une architecture MVC moderne avec separation claire des responsabilites.'
)
add_bullet_list([
    'Backend : Laravel 11 avec Actions pattern (CreateJobAction, MoveApplicationStageAction, etc.)',
    'Frontend : Vue 3 avec Composition API + Inertia.js pour le SPA-like experience',
    'Multi-tenancy : Scope automatique par company_id via le trait BelongsToCompany',
    'Middleware : SetCurrentCompany, EnsureCompanySubscription, CheckUsageLimits',
    'RBAC : 5 roles (owner, admin, recruiter, hiring_manager, interviewer) via Spatie Permissions',
    'Jobs asynchrones : ParseResumeJob, ScoreCandidateJob, SummarizeApplicationJob',
])

# ─── Module 1: Authentication ───
add_section_page(
    'Module 1 : Authentification',
    'Systeme d\'authentification complet base sur Laravel Breeze avec inscription, '
    'connexion, et gestion du profil.'
)
add_screenshot('11-login.png', 'Figure 1 : Page de connexion')
add_screenshot('12-register.png', 'Figure 2 : Page d\'inscription')

# ─── Module 2: Dashboard ───
add_section_page(
    'Module 2 : Tableau de Bord',
    'Le tableau de bord offre une vue d\'ensemble des KPIs cles du recrutement : '
    'offres ouvertes, candidats actifs, entretiens du jour, et delai moyen d\'embauche.'
)
add_screenshot('01-dashboard.png', 'Figure 3 : Tableau de bord avec KPIs et activite recente')

# ─── Module 3: Job Postings ───
add_section_page(
    'Module 3 : Gestion des Offres d\'Emploi',
    'CRUD complet pour les offres d\'emploi avec filtrage par statut, departement et localisation. '
    'Chaque offre peut etre configuree avec des competences requises, salaire, type de contrat, etc.'
)
add_screenshot('02-jobs-list.png', 'Figure 4 : Liste des offres d\'emploi avec filtres')
add_screenshot('03-jobs-create.png', 'Figure 5 : Formulaire de creation d\'une offre d\'emploi')

# ─── Module 4: Candidates ───
add_section_page(
    'Module 4 : Gestion des Candidats',
    'Base de donnees centralisee des candidats avec profil complet incluant '
    'experiences, formations, competences et resume. Support pour l\'upload et le parsing de CV.'
)
add_screenshot('04-candidates-list.png', 'Figure 6 : Liste des candidats avec recherche')
add_screenshot('05-candidates-create.png', 'Figure 7 : Formulaire d\'ajout d\'un candidat')

# ─── Module 5: Interviews ───
add_section_page(
    'Module 5 : Planification des Entretiens',
    'Systeme de planification d\'entretiens avec support pour differents types '
    '(telephonique, video, sur site, technique, RH). Filtrage par statut et plage de dates.'
)
add_screenshot('06-interviews.png', 'Figure 8 : Liste des entretiens planifies')

# ─── Module 6: Offers ───
add_section_page(
    'Module 6 : Gestion des Offres',
    'Suivi des offres envoyees aux candidats avec salaire, statut, et dates d\'envoi/reponse. '
    'Les candidats peuvent accepter ou decliner via un lien signe.'
)
add_screenshot('07-offers.png', 'Figure 9 : Liste des offres envoyees')

# ─── Module 7: Analytics ───
add_section_page(
    'Module 7 : Analytique',
    'Tableaux de bord analytiques avec KPIs, entonnoir du pipeline de recrutement, '
    'et tendances du delai d\'embauche.'
)
add_screenshot('08-analytics.png', 'Figure 10 : Dashboard analytique avec KPIs et graphiques')

# ─── Module 8: Profile ───
add_section_page(
    'Module 8 : Profil Utilisateur',
    'Gestion du profil utilisateur avec modification du nom, email et mot de passe.'
)
add_screenshot('09-profile.png', 'Figure 11 : Page de gestion du profil')

# ─── Module 9: Career Portal ───
add_section_page(
    'Module 9 : Portail Carrieres Public',
    'Page publique de carrieres personnalisee par entreprise, accessible via un slug unique. '
    'Les visiteurs peuvent rechercher et filtrer les postes ouverts, puis postuler directement.'
)
add_screenshot('10-careers.png', 'Figure 12 : Portail carrieres public')

# ─── AI Features ───
add_section_page(
    'Module 10 : Fonctionnalites IA',
    'RecruitAI integre l\'intelligence artificielle via OpenAI pour automatiser '
    'les taches repetitives du recrutement.'
)
add_bullet_list([
    'Parsing de CV : Extraction automatique des informations (nom, email, competences, experience)',
    'Scoring des candidats : Evaluation automatique de l\'adequation candidat/poste',
    'Resume de candidature : Generation automatique d\'un resume pour les recruteurs',
    'Suggestions de questions : Generation de questions d\'entretien basees sur le profil',
    'Enrichissement de description : Amelioration automatique des descriptions de poste',
])

# ─── Billing ───
add_section_page(
    'Module 11 : Facturation',
    'Systeme de facturation integre via Stripe (Laravel Cashier) avec trois plans tarifaires.'
)
add_bullet_list([
    'Plan Free : Fonctionnalites de base avec limites d\'utilisation',
    'Plan Starter : Fonctionnalites etendues pour les petites equipes',
    'Plan Professional : Acces complet a toutes les fonctionnalites',
    'Gestion des abonnements, moyens de paiement et factures',
    'Middleware CheckUsageLimits pour le controle des quotas',
])

# ─── Pipeline ───
add_section_page(
    'Module 12 : Pipeline de Recrutement',
    'Vue Kanban du pipeline de recrutement permettant de suivre les candidatures '
    'a travers les differentes etapes du processus.'
)
add_bullet_list([
    'Drag & drop des candidatures entre les etapes',
    'Etapes configurables (Nouveau, Pre-selection, Entretien, Offre, Embauche)',
    'Rejet avec motif optionnel',
    'Historique des mouvements',
])

# ─── Conclusion ───
add_section_page(
    'Conclusion',
    ''
)
add_bullet_list([
    'Application complete couvrant l\'ensemble du cycle de recrutement',
    'Interface intuitive et entierement en francais',
    'Architecture moderne et maintenable (Laravel 11 + Vue 3)',
    'Integration IA pour l\'automatisation des taches repetitives',
    'Multi-tenancy et RBAC pour la securite et l\'isolation des donnees',
    'Systeme de facturation Stripe pour la monetisation',
])

doc.save(OUTPUT_PATH)
print(f'Presentation saved to: {OUTPUT_PATH}')
print(f'File size: {os.path.getsize(OUTPUT_PATH) / 1024:.0f} KB')
